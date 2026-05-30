"""
DOCX converter for MarkItDown-Pro

Note: This converter requires python-docx to be installed.
Install with: pip install python-docx
"""

from typing import Dict, Any, Optional
from .base import BaseConverter


class DOCXConverter(BaseConverter):
    """
    Converter for DOCX files
    
    Converts Word documents to Markdown with formatting preservation.
    Requires python-docx: pip install python-docx
    """
    
    name = "docx"
    description = "DOCX to Markdown converter"
    supported_extensions = ['.docx']
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize DOCX converter
        
        Args:
            config: Optional configuration
        """
        super().__init__(config)
        
        # Try to import docx
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX conversion. "
                "Install with: pip install python-docx"
            )
    
    def convert(self, file_path: str, options: Optional[Dict[str, Any]] = None) -> str:
        """
        Convert DOCX file to Markdown
        
        Args:
            file_path: Path to DOCX file
            options: Conversion options
            
        Returns:
            Markdown content
        """
        options = options or {}
        include_tables = options.get('include_tables', True)
        include_images = options.get('include_images', False)
        
        lines = []
        
        try:
            doc = self.docx.Document(file_path)
            
            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    lines.append("")
                    continue
                
                # Determine paragraph style
                style_name = para.style.name if para.style else "Normal"
                
                # Convert based on style
                if 'Heading' in style_name:
                    level = self._get_heading_level(style_name)
                    lines.append(f"{'#' * level} {para.text}")
                elif 'Title' in style_name:
                    lines.append(f"# {para.text}")
                else:
                    # Regular paragraph with formatting
                    text = self._process_paragraph(para)
                    lines.append(text)
            
            # Process tables
            if include_tables:
                for table in doc.tables:
                    lines.append("")
                    lines.append(self._convert_table(table))
                    lines.append("")
            
        except Exception as e:
            lines.append(f"**Error converting DOCX**: {str(e)}")
        
        return '\n'.join(lines)
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        Get heading level from style name
        
        Args:
            style_name: Style name
            
        Returns:
            Heading level (1-6)
        """
        import re
        
        match = re.search(r'Heading\s*(\d+)', style_name)
        if match:
            level = int(match.group(1))
            return max(1, min(6, level))
        return 1
    
    def _process_paragraph(self, para) -> str:
        """
        Process paragraph with formatting
        
        Args:
            para: Paragraph object
            
        Returns:
            Formatted text
        """
        parts = []
        
        for run in para.runs:
            text = run.text
            
            # Apply formatting
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _convert_table(self, table) -> str:
        """
        Convert table to Markdown
        
        Args:
            table: Table object
            
        Returns:
            Markdown table
        """
        rows = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get text from cell
                text = cell.text.strip().replace('|', '\\|')
                row_data.append(text)
            rows.append(row_data)
        
        if not rows:
            return ""
        
        # Build Markdown table
        lines = []
        
        # Header row
        headers = rows[0]
        lines.append('| ' + ' | '.join(headers) + ' |')
        lines.append('|' + '|'.join(['---' for _ in headers]) + '|')
        
        # Data rows
        for row in rows[1:]:
            lines.append('| ' + ' | '.join(row) + ' |')
        
        return '\n'.join(lines)
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary of metadata
        """
        metadata = super().get_metadata(file_path)
        
        try:
            doc = self.docx.Document(file_path)
            
            # Count elements
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # Extract core properties
            if doc.core_properties:
                props = doc.core_properties
                if props.title:
                    metadata['title'] = props.title
                if props.author:
                    metadata['author'] = props.author
                if props.subject:
                    metadata['subject'] = props.subject
                if props.created:
                    metadata['created'] = props.created.isoformat() if hasattr(props.created, 'isoformat') else str(props.created)
                if props.modified:
                    metadata['modified'] = props.modified.isoformat() if hasattr(props.modified, 'isoformat') else str(props.modified)
            
        except Exception:
            pass
        
        return metadata
